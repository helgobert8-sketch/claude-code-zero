"""Minimal, dependency-light Markdown -> .docx converter (python-docx).
Handles: ATX headings, GFM pipe tables, fenced code blocks, ordered/unordered
lists, blockquotes, horizontal rules, inline **bold** and `code`.
Usage: py md_to_docx.py <input.md> <output.docx>
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(paragraph, text):
    """Add runs to a paragraph, parsing **bold** and `code`."""
    # split on bold and inline-code tokens, keeping delimiters
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(tok)


def main():
    src, dst = sys.argv[1], sys.argv[2]
    with open(src, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.lstrip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].lstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # GFM pipe table
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for j, htext in enumerate(header):
                hdr[j].paragraphs[0].text = ""
                r = hdr[j].paragraphs[0].add_run(htext.replace("**", ""))
                r.bold = True
                r.font.size = Pt(10)
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    val = row[j] if j < len(row) else ""
                    cells[j].paragraphs[0].text = ""
                    add_inline(cells[j].paragraphs[0], val)
                    for rr in cells[j].paragraphs[0].runs:
                        rr.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        # horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            p = doc.add_paragraph()
            p.add_run("_" * 50).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            h.text = ""
            add_inline(h, text)
            i += 1
            continue

        # blockquote
        if line.lstrip().startswith(">"):
            text = re.sub(r"^\s*>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run("")
            add_inline(p, text)
            for rr in p.runs:
                rr.italic = True
                rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m and not line.strip().startswith("---"):
            text = m.group(2)
            # render task-list checkboxes as visible glyphs
            text = re.sub(r"^\[ \]\s*", "☐ ", text)
            text = re.sub(r"^\[[xX]\]\s*", "☑ ", text)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
            i += 1
            continue

        # ordered list
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(1))
            i += 1
            continue

        # blank line
        if line.strip() == "":
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(dst)
    print(f"OK: wrote {dst}")


if __name__ == "__main__":
    main()
